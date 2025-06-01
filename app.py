# Initial Push nach GitHub

from flask import Flask, request, jsonify, render_template, send_file, url_for
from flask_cors import CORS
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
import openai
import os

# === Initialisierung ===
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

# === Routen ===

@app.route('/')
def index():
    return render_template('index.html')  # muss in templates/index.html liegen

@app.route('/api/generate', methods=['POST'])
def generate():
    data = request.get_json()
    inhalt = data.get('inhalt', '').strip()
    stil = data.get('stil', 'modern')

    if not inhalt:
        return jsonify({'error': 'Keine Inhalte übergeben'}), 400

    prompt = f"""
Du bist ein professioneller Bewerbungscoach für den deutschsprachigen Raum.
Erstelle basierend auf folgendem Input eine überzeugende Bewerbung im Stil: {stil}

Input:
{inhalt}

Antwortformat:
Anschreiben:
[Text]

Lebenslauf:
[Text]
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Du bist ein professioneller deutscher Bewerbungscoach."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )

        full_text = response.choices[0].message['content']
        if "Lebenslauf:" in full_text:
            anschreiben, lebenslauf = full_text.split("Lebenslauf:", 1)
        else:
            anschreiben = full_text
            lebenslauf = "Keine klare Trennung gefunden."

        return jsonify({
            'anschreiben': anschreiben.replace("Anschreiben:", "").strip(),
            'lebenslauf': lebenslauf.strip()
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/docx', methods=['POST'])
def export_docx():
    data = request.get_json()
    anschreiben = data.get('anschreiben', '')
    lebenslauf = data.get('lebenslauf', '')

    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        doc.add_heading('Anschreiben', level=1)
        for line in anschreiben.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        doc.add_page_break()
        doc.add_heading('Lebenslauf', level=1)
        for line in lebenslauf.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="bewerbung.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    data = request.get_json()
    anschreiben = data.get('anschreiben', '')
    lebenslauf = data.get('lebenslauf', '')

    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm
        )

        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("Anschreiben", styles['Heading1']))
        for line in anschreiben.split('\n'):
            if line.strip():
                story.append(Paragraph(line.strip(), styles['Normal']))
                story.append(Spacer(1, 6))

        story.append(PageBreak())
        story.append(Paragraph("Lebenslauf", styles['Heading1']))
        for line in lebenslauf.split('\n'):
            if line.strip():
                story.append(Paragraph(line.strip(), styles['Normal']))
                story.append(Spacer(1, 6))

        doc.build(story)

        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="bewerbung.pdf",
            mimetype="application/pdf"
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# === Starten ===
if __name__ == '__main__':
    app.run(debug=True)

if __name__ == '__main__':
    import os
    port = int(os.environ.get("PORT", 5000))  # <-- PORT von Render
    app.run(host='0.0.0.0', port=port, debug=True)
